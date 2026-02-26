import os
import time
import logging
import requests
import uuid
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from dotenv import load_dotenv
from docx2pdf import convert
from PIL import Image
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from pptx import Presentation
from pptx.util import Inches as PPTInches, Pt as PPTPt
from pptx.enum.text import PP_PARAGRAPH_ALIGNMENT

# ── Load env ──
load_dotenv()
PORT = int(os.getenv("PORT", 5002))

# ── Logging ──
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s"
)
log = logging.getLogger(__name__)

# ── Flask & CORS ──
app = Flask(__name__)
CORS(app)

# ── Directories ──
BASE_DIR   = os.path.dirname(__file__)
EXPORT_DIR = os.path.join(BASE_DIR, "exports")
DIAGR_DIR  = os.path.join(EXPORT_DIR, "diagrams")
os.makedirs(EXPORT_DIR, exist_ok=True)
os.makedirs(DIAGR_DIR, exist_ok=True)

# ── Helpers ──
def sanitize_image(path):
    try:
        with Image.open(path) as img:
            rgb = img.convert("RGB")
            directory, filename = os.path.split(path)
            name, ext = os.path.splitext(filename)
            safe_filename = f"{name}_safe.png"
            safe_path = os.path.join(directory, safe_filename)
            rgb.save(safe_path, format="PNG", optimize=True)
            return safe_path
    except Exception as e:
        log.warning("Failed to sanitize %s: %s", path, e)
        return None

def set_shading_for_paragraph(paragraph, fill_color):
    p_element = paragraph._p
    p_pr = p_element.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    p_pr.append(shd)

# ── Proxy to UML Agent ──
@app.route('/generate-uml', methods=['POST'])
def proxy_uml():
    # This endpoint is kept for potential direct testing but is not used in the main flow.
    payload = request.get_json(force=True)
    resp = requests.post(
        os.getenv("UML_AGENT_URL", "https://umlgenerator.onrender.com") + "/generate-uml-image",
        json={
            "build_id":     payload.get("build_id", "default"),
            "abstract":     payload.get("abstract", ""),
            "instructions": payload.get("instructions", "")
        }, timeout=120
    )
    return jsonify(resp.json()), resp.status_code

# ── Ingest UML diagrams ──
@app.route('/ingest-diagram', methods=['POST'])
def ingest_diagram():
    build_id = request.form.get('build_id')
    if not build_id:
        return jsonify({"error": "build_id is required"}), 400

    diagramType = request.form.get('diagramType', 'diagram')
    index       = request.form.get('index', '1')
    img_file    = request.files.get('image')
    if not img_file:
        return jsonify({"error": "no image"}), 400

    target_dir = os.path.join(DIAGR_DIR, build_id)
    os.makedirs(target_dir, exist_ok=True)
    filename = f"{diagramType}_{index}.png"
    path     = os.path.join(target_dir, filename)
    img_file.save(path)
    log.info("Ingested diagram for build [%s] -> %s", build_id, path)
    return jsonify({"path": path}), 200

# ── Build document ──
@app.route('/build-document', methods=['POST'])
def build_document():
    start = time.time()
    data  = request.get_json(force=True)
    build_id = str(uuid.uuid4())

    # 1) Get Markdown
    log.info("Requesting markdown from parent agent for build [%s]", build_id)
    try:
        md_resp = requests.post(
            os.getenv("PARENT_AGENT_URL", "https://aiagent-xyq4.onrender.com") + "/generate-doc",
            json={
                "code":          data.get("code"),
                "project_info":  data.get("project_info"),
                "instructions":  data.get("instructions"),
                "pages":         data.get("pages", 1),
                "return_format": "markdown"
            }, timeout=120
        )
        md_resp.raise_for_status()
        raw_md = md_resp.text or ""
        if not raw_md.strip():
            return jsonify({"error":"Empty markdown response"}), 500
    except Exception as e:
        log.error("Failed to get markdown: %s", e)
        return jsonify({"error": "Failed to generate document content"}), 500

    # 2) Get UML diagrams
    log.info("Requesting diagrams from UML agent for build [%s]", build_id)
    build_diagrams_dir = os.path.join(DIAGR_DIR, build_id)
    os.makedirs(build_diagrams_dir, exist_ok=True)
    try:
        uml_resp = requests.post(
            os.getenv("UML_AGENT_URL", "https://umlgenerator.onrender.com") + "/generate-uml-image",
            json={
                "build_id":     build_id,
                "abstract":     data.get("abstract",""),
                "instructions": data.get("uml_instructions","")
            }, timeout=120
        )
        uml_resp.raise_for_status()
        diagram_specs = uml_resp.json().get("diagrams", [])
        log.info("Got %d diagram specifications", len(diagram_specs))
    except Exception as e:
        log.error("Failed to get diagrams: %s", e)
        diagram_specs = []
 
    # 3) Build Polished DOCX
    doc = Document()

    # Define Styles
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Poppins'
    font.size = Pt(11)

    style_h1 = doc.styles['Heading 1']
    style_h1.font.name = 'Poppins'
    style_h1.font.size = Pt(18)
    style_h1.font.bold = True
    style_h1.font.color.rgb = RGBColor(0x3B, 0x4B, 0x64)
    style_h1.paragraph_format.space_after = Pt(12)

    style_h2 = doc.styles['Heading 2']
    style_h2.font.name = 'Poppins'
    style_h2.font.size = Pt(14)
    style_h2.font.bold = True
    style_h2.paragraph_format.space_after = Pt(10)

    # Add Cover Page
    project_title = data.get("instructions", "AI-Generated Documentation").split('\n')[0]
    cover_title = doc.add_heading('Technical Report', level=0)
    cover_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p = doc.add_paragraph()
    p.add_run(project_title).bold = True
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph(f"Generated on: {time.strftime('%B %d, %Y')}", style='Caption').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_page_break()

    # Add Footer with Page Numbers
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.text = " DocuAgent | Confidential\t"
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # Process Markdown into Document Body
    is_code_block = False
    for line in raw_md.splitlines():
        clean_line = line.strip()

        if '```' in clean_line:
            is_code_block = not is_code_block
            continue

        if is_code_block:
            p = doc.add_paragraph()
            run = p.add_run(clean_line)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            set_shading_for_paragraph(p, "F1F1F1")
            p.paragraph_format.space_after = Pt(4)
        elif clean_line.startswith('# '):
            doc.add_heading(clean_line[2:].strip(), level=1)
        elif clean_line.startswith('## '):
            doc.add_heading(clean_line[3:].strip(), level=2)
        elif clean_line.startswith('### '):
            doc.add_heading(clean_line[4:].strip(), level=3)
        elif clean_line:
            p = doc.add_paragraph(clean_line.replace('*', ''))
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            p.paragraph_format.space_after = Pt(8)

    # Add Diagrams
    if diagram_specs:
        doc.add_heading("Diagrams", level=1)

    for spec in diagram_specs:
        diagram_type_raw = spec.get('diagramType', 'Diagram')
        description = spec.get('description', '')
        title = diagram_type_raw.replace('_', ' ').replace('-', ' ').title()
        doc.add_heading(title, level=2)

        p = doc.add_paragraph(description)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        p.paragraph_format.space_after = Pt(8)

        img_filename = f"{diagram_type_raw}_{spec.get('index', 1)}.png"
        img_path = os.path.join(build_diagrams_dir, img_filename)

        if os.path.isfile(img_path):
            safe_path = sanitize_image(img_path)
            if safe_path:
                try:
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    p_img.add_run().add_picture(safe_path, width=Inches(5.5))
                except Exception as e:
                    log.warning("Embed %s failed: %s", safe_path, e)
        else:
            log.warning("Missing image for spec: %s", img_path)
        doc.add_paragraph()

    # Save Files
    ts        = int(time.time())
    docx_file = f"combined_{ts}.docx"
    pdf_file  = f"combined_{ts}.pdf"
    pptx_file = f"combined_{ts}.pptx"

    docx_path = os.path.join(EXPORT_DIR, docx_file)
    pdf_path  = os.path.join(EXPORT_DIR, pdf_file)
    pptx_path = os.path.join(EXPORT_DIR, pptx_file)

    try:
        doc.save(docx_path)
        log.info("DOCX saved → %s (%d bytes)", docx_path, os.path.getsize(docx_path))
    except Exception as e:
        log.error("DOCX save failed: %s", e)
        return jsonify({"error":"DOCX save failed"}), 500

    if os.path.exists(docx_path) and os.path.getsize(docx_path) > 1000:
        try:
            convert(docx_path, pdf_path)
            log.info("PDF saved → %s", pdf_path)
        except Exception as e:
            log.error("PDF conversion failed: %s", e)

    # Note: PPTX generation is not implemented yet
    prs = Presentation()
    try:
        prs.save(pptx_path)
        log.info("PPTX (empty) saved → %s", pptx_path)
    except Exception as e:
        log.error("PPTX failed: %s", e)

    log.info("Completed build %s in %.2fs", build_id, time.time() - start)

    response_data = { "diagrams_count": len(diagram_specs) }
    if os.path.exists(docx_path): response_data["docx"] = docx_file
    if os.path.exists(pdf_path): response_data["pdf"] = pdf_file
    if os.path.exists(pptx_path): response_data["pptx"] = pptx_file

    return jsonify(response_data), 200

# ── Download Generated Files ──
@app.route('/download/<filetype>/<filename>', methods=['GET'])
def download_file(filetype, filename):
    full_path = os.path.join(EXPORT_DIR, filename)
    if not os.path.isfile(full_path):
        return jsonify({"error":"File not found on server"}), 404
    return send_file(full_path, as_attachment=True, download_name=filename)



if __name__ == "__main__":
    app.run(host="0.0.0.0", port=PORT, debug=True)
