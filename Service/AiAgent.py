# Aigenerator Service
import os
import logging
import ast
from io import BytesIO
from datetime import datetime

from flask import Flask, request, jsonify, send_file
from flask_cors import CORS

import markdown
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from dotenv import load_dotenv
import google.generativeai as genai

# ── Configuration ──
load_dotenv()
logging.basicConfig(level=logging.INFO, format='%(asctime)s %(levelname)s: %(message)s')

app = Flask(__name__)
CORS(app)

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
if not GEMINI_API_KEY:
    raise RuntimeError("GEMINI_API_KEY not set")
genai.configure(api_key=GEMINI_API_KEY)

# ── Default Prompt for Agent ──
DEFAULT_INSTRUCTIONS = """
You are an expert technical documentation agent tasked with generating professional, concise, and academically styled project reports for software engineering deliverables (e.g., AI systems, mobile apps, web applications).

For **each** section below, produce **3–5 formal paragraphs** (at least **15 lines**) in clear academic English. Use **bolded** section titles, **numbered** subheadings, consistent indentation, and realistic examples or technical explanations wherever relevant. Ensure coherence, clarity, and logical flow across sections.

### Table of Contents
1. **Introduction**
 1.1 Motivation
 1.2 Problem Definition
 1.3 Objectives
 1.4 Scope & Limitations

2. **Literature Survey**
 2.1 Overview of Related Work
 2.2 Existing Systems & Their Drawbacks
 2.3 Proposed System & Innovations
 2.4 Feasibility Analysis
 2.5 Key Features
 2.6 Required Technologies

3. **Analysis**
 3.1 Analytical Overview
 3.2 Requirements Specification
  3.2.1 Functional Requirements
  3.2.2 Non-functional Requirements
  3.2.3 Software & Hardware Specifications

4. **Design**
 4.1 Architectural Overview
 4.2 System Architecture Diagram
 4.4 Module Organization & Interfaces

5. **Implementation**
 5.1 Implementation Strategy
 5.2 Core Functional Modules
 5.3 Representative Code Snippets
 5.4 Deployment & Execution
  5.4.1 Results Analysis
  5.4.2 Sample Output Screens

6. **Testing, Validation & Results**
 6.1 Testing Overview
 6.2 Test Methodologies
 6.3 Test Case Design & Scenarios
 6.4 Validation of Outcomes

7. **Conclusion & Future Work**
 7.1 Summary of Achievements
 7.2 Limitations & Lessons Learned
 7.3 Recommendations for Future Enhancements
"""

# ── Code Analyzer ──
def parse_code(code: str, ext: str):
    logging.info(f"Parsing code with extension '{ext}'")
    if ext == ".py":
        try:
            tree = ast.parse(code)
        except SyntaxError:
            logging.error("Invalid Python syntax detected")
            return {"error": "Invalid Python syntax"}
        funcs, classes = [], []
        for node in ast.walk(tree):
            if isinstance(node, ast.FunctionDef):
                funcs.append(node.name)
            elif isinstance(node, ast.ClassDef):
                classes.append(node.name)
        return {"functions": funcs, "classes": classes, "lines": len(code.splitlines())}
    return {"functions": [], "classes": [], "lines": len(code.splitlines())}

# ── Gemini Agent ──
def call_gemini(code: str, project_info: str, instructions: str, pages: int = 1):
    logging.info(f"Calling Gemini AI to generate ~{pages} page(s) of documentation")

    full_prompt = f"""
Use the following to generate a **Markdown** technical report of approximately {pages} page{'s' if pages > 1 else ''}:

**Project Description:**
\"\"\"{project_info}\"\"\"

**User Instructions:**
\"\"\"{instructions.strip()}\"\"\"

**Source Code (for context):**
\"\"\"{code}\"\"\"

Produce only the final well-structured Markdown text with bolded headings, numbered sections, and consistent formatting.
"""
    model = genai.GenerativeModel("gemini-2.5-flash")
    response = model.generate_content(full_prompt)
    return response.text.strip()

# ── Output Generators ──
def generate_pdf_from_text(text: str):
    logging.info("Generating PDF from Markdown text")
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    x, y = 40, height - 40
    for line in text.splitlines():
        if y < 50:
            c.showPage()
            y = height - 40
        c.drawString(x, y, line.strip())
        y -= 15
    c.save()
    buffer.seek(0)
    return buffer

def generate_docx_from_text(text: str):
    logging.info("Generating DOCX from Markdown text")
    doc = Document()
    for line in text.splitlines():
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        else:
            doc.add_paragraph(line)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ── API Endpoint ──
@app.route("/generate-doc", methods=["POST"])
def generate_doc():
    data = request.get_json()
    logging.info(f"Received POST data keys: {list(data.keys())}")

    file_path     = data.get("file_path")
    code          = data.get("code")
    extension     = data.get("extension", ".txt").lower()
    project_info  = data.get("project_info")
    instructions  = data.get("instructions", DEFAULT_INSTRUCTIONS)
    return_format = data.get("return_format", "pdf").lower()
    pages         = data.get("pages", 1)

    try:
        pages = max(1, int(pages))
    except Exception:
        pages = 1

    if file_path:
        if not os.path.isfile(file_path):
            logging.error(f"File not found: {file_path}")
            return jsonify({"error": "File not found"}), 400
        with open(file_path, "r", encoding="utf-8") as f:
            code = f.read()
        extension = os.path.splitext(file_path)[1].lower()
        logging.info(f"Loaded code from: {file_path}")

    if not code or not project_info:
        logging.error("Missing required fields: 'code' or 'project_info'")
        return jsonify({"error": "Missing code or project_info"}), 400

    parse_info = parse_code(code, extension)
    if "error" in parse_info:
        return jsonify(parse_info), 400

    markdown_text = call_gemini(code, project_info, instructions, pages)

    if return_format == "markdown":
        return (
            markdown_text,
            200,
            {
                "Content-Type": "text/markdown",
                "Content-Disposition": "attachment; filename=documentation.md"
            }
        )
    elif return_format == "pdf":
        return send_file(
            generate_pdf_from_text(markdown_text),
            mimetype="application/pdf",
            as_attachment=True,
            download_name="documentation.pdf"
        )
    elif return_format == "docx":
        return send_file(
            generate_docx_from_text(markdown_text),
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name="documentation.docx"
        )
    elif return_format == "text":
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"documentation_{timestamp}.txt"
        save_dir = "saved_docs"
        os.makedirs(save_dir, exist_ok=True)
        file_path = os.path.join(save_dir, filename)
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(markdown_text)
        return send_file(
            file_path,
            mimetype="text/plain",
            as_attachment=True,
            download_name=filename
        )
    else:
        logging.error(f"Unsupported return format: {return_format}")
        return jsonify({"error": f"Unsupported format: {return_format}"}), 400

if __name__ == "__main__":
    port = int(os.environ["PORT"])
    app.run(host="0.0.0.0", port=port, debug=False)
